## after fill out csv and come back run modify_para to create final output cover letters

def modify_para(Templatedoc, csvfilename, outputfiledocPath):

    import docx
    import pandas as pd  
    document = docx.Document() 
    doc1 = docx.Document(Templatedoc)
    for num in range(len(doc1.paragraphs)):


        doc = docx.Document(Templatedoc)

        target_para = doc.paragraphs[num].text
        #print(target_para)
        df = pd.read_csv(csvfilename)
        dictt = {}

        for index, row in df.iterrows():
            dictt[row['tag']] = row['Input']
        # print(dictt)
        keywordList = {}
        for key, val in dictt.items():

            if key in doc.paragraphs[num].text:
                keywordList[key] = val

        countkey = len(keywordList)
        # print(keywordList)

        for key, val in keywordList.items():

                if key in target_para:
                    change_para = target_para.replace(key, val.strip(" "))
                    target_para = change_para 

        print(target_para)
        
        

        document.add_paragraph(target_para, style='NoSpacing')


    #     print(target_para)


    document.save(outputfiledocPath)

    
    