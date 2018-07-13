def createCSV(Templatedoc, csvfilename):

        import pandas as pd
        import docx
        def getText(docfilename):
            
            doc = docx.Document(docfilename)
            fullText = []
            for para in doc.paragraphs:
                fullText.append(para.text)
            return '\n'.join(fullText)

        data= getText(Templatedoc).split()
        infolist = []
        itemlist = []
        namelist = []
        for item in data:
            if item.startswith('<'):
                namelist.append(item)
                itemname = item.strip("<").strip(">").strip(">.").strip(">:").strip('>-')
                itemlist.append(itemname)
        print(itemlist)
        print(namelist)
        doc = docx.Document(Templatedoc)
        print('Parapraph length: ',  len(doc.paragraphs))

        d = pd.DataFrame(namelist, itemlist)

        d = d.reset_index()
        d= d.rename(columns = {"": "name", 0:"tag"})
        d.insert(loc = 2, column='Input', value = " ")
        d.to_csv(csvfilename, index= False)


